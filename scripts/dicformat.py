import docx

def format_words(output_file, sorted_list):
    """
    Coloring and formating to docx
    """
    
    doc = docx.Document()
    i = 0

    for w in sorted_list:
        temp = w.split(":")
        para = doc.add_paragraph()

        if len(temp) > 1:
            i += 1
            word = para.add_run(u'%s : ' % temp[0])
            word.font.size = docx.shared.Pt(12)
            word.font.name = "Calibri"
            word.font.color.rgb = docx.shared.RGBColor(0, 176, 80) #The color, YAAAY!

            definition = para.add_run(u'%s' % temp[1])
            definition.font.size = docx.shared.Pt(12)
            definition.font.name = "Calibri"
            doc.save(output_file + ".docx")
        else:
            # If it's an header eg: A(a)
            word = para.add_run(u'%s' % temp[0])
            word.font.size = docx.shared.Pt(24)
            word.font.name = "Cooper Std Black"
            word.font.color.rgb = docx.shared.RGBColor(255, 0, 0)

    with open("nbMots.txt", "w") as file:
        file.write("Nombre de mots: %s" % i)

    print("Nombre de mots: %s" % i)

